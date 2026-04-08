#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gerador da Seção Metodologia v.1 — ICONOCRACY
Formatação ABNT NBR 14724: Times New Roman 12pt, margens 3/2/3/2 cm, espaço 1,5
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = "/Users/ana/iconocracy-corpus/PHD/Metodologia_v1_ICONOCRACY.docx"

doc = Document()

# ── Margens ABNT: sup 3cm, inf 2cm, esq 3cm, dir 2cm ──────────────────────
for section in doc.sections:
    section.top_margin    = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)

# ── Estilo Normal: Times New Roman 12pt ────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)

def abnt_paragraph(doc, text="", bold=False, italic=False, size=12,
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_cm=1.25,
                   space_before=0, space_after=200, spacing=1.5, color=None):
    """Parágrafo ABNT com espaçamento 1,5 e recuo de 1,25 cm."""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.first_line_indent = Cm(indent_cm)
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after / 20)   # Pt usa 1/12750 EMU; passamos direto
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = spacing
    if text:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.bold   = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def heading(doc, text, level=1):
    """Título de seção ABNT: maiúsculas/negrito para H1, negrito para H2."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after  = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0)
    run = p.add_run(text.upper() if level == 1 else text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    return p

def separator(doc):
    abnt_paragraph(doc, "", space_before=0, space_after=0)

# ══════════════════════════════════════════════════════════════════════
# CAPA / CABEÇALHO
# ══════════════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run("METODOLOGIA (v.1)")
r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run("ICONOCRACIA: Alegoria Feminina na História da Cultura Jurídica (Séculos XIX–XX)")
r2.font.name = 'Times New Roman'; r2.font.size = Pt(12); r2.bold = True

abnt_paragraph(doc, "Doutoranda: Ana Vanzin | Programa de Pós-Graduação em Direito — PPGD/UFSC",
               align=WD_ALIGN_PARAGRAPH.CENTER, indent_cm=0, space_after=0, bold=False)
abnt_paragraph(doc, "Disciplina: Métodos e Metodologias Aplicados ao Projeto de Tese | Encontro 5 · 09/04/2026",
               align=WD_ALIGN_PARAGRAPH.CENTER, indent_cm=0, space_after=12)

# ══════════════════════════════════════════════════════════════════════
# 1. POSICIONAMENTO EPISTEMOLÓGICO
# ══════════════════════════════════════════════════════════════════════
heading(doc, "1  Posicionamento Epistemológico e Tradição de Pesquisa")

abnt_paragraph(doc,
    "A pesquisa se inscreve em uma tradição interdisciplinar que articula história do direito, "
    "história da arte e teoria de gênero. Do ponto de vista epistemológico, adota-se uma perspectiva "
    "qualitativa, histórico-interpretativa, que reconhece a imagem como documento histórico dotado de "
    "sentido jurídico-político próprio. A abordagem não é positivista — não busca leis gerais a partir "
    "de dados quantitativos —, mas tampouco é puramente hermenêutica no sentido diltheyano: parte-se "
    "da premissa de que as alegorias femininas, longe de serem meros ornamentos visuais, constituem "
    "dispositivos normativos que participam ativamente da construção e legitimação da ordem jurídica estatal.")

abnt_paragraph(doc,
    "A tradição de pesquisa em que se insere esta tese é a da história da cultura jurídica visual, "
    "campo emergente que combina o historicismo jurídico humanista cultivado por Paolo Grossi, "
    "António Manuel Hespanha e Arno Dal Ri Jr. com os métodos da história da arte — especialmente "
    "o método iconológico de Erwin Panofsky e a noção warburguiana de Pathosformel — e com a teoria "
    "feminista do contrato social (Carole Pateman) e da produção simbólica de gênero (Marina Warner, "
    "Maurice Agulhon). A interdisciplinaridade não é ornamental: ela é constitutiva do objeto de "
    "pesquisa, pois a alegoria feminina da Justiça ou da República só pode ser adequadamente "
    "compreendida na interseção entre o campo jurídico, o campo visual e o campo das relações de gênero.")

abnt_paragraph(doc,
    "Nessa perspectiva, a imagem não é mero reflexo passivo das estruturas jurídicas: ela as produz, "
    "as naturaliza e as contesta. A cultura jurídica visual designa o conjunto de práticas, objetos e "
    "saberes pelos quais as sociedades modernas constituem e reproduzem, por meios visuais, sua ordem "
    "jurídica imaginada. A alegoria feminina estatal — seja a Marianne francesa, a Britannia britânica, "
    "a Germania wilhelminiana, a Columbia norte-americana, a Belgique constitucional ou a República "
    "brasileira positivista — é o operador visual central desse processo.")

# ══════════════════════════════════════════════════════════════════════
# 2. NATUREZA E TIPO DE PESQUISA
# ══════════════════════════════════════════════════════════════════════
heading(doc, "2  Natureza e Tipo de Pesquisa")

abnt_paragraph(doc,
    "A pesquisa é de natureza qualitativa, exploratória e histórico-comparada. É exploratória porque "
    "o campo da iconologia jurídica aplicada ao gênero — especialmente no âmbito do direito público "
    "comparado — carece de sistematização: os estudos existentes, como os de Stefan Huygebaert, "
    "Georges Martyn, Bettina Kocher e Gerlinde Schott, são majoritariamente monográficos e não "
    "estabelecem um protocolo de análise comparada entre múltiplas jurisdições nacionais ao longo "
    "de dois séculos. É histórica porque o corpus se estende por duzentos anos (1800–2000) e "
    "demanda contextualização em processos políticos de longa duração: fundações republicanas, "
    "consolidações institucionais, ciclos de guerra e imperialismo, descolonizações. É comparada "
    "porque examina seis jurisdições nacionais de modo sistemático, identificando convergências "
    "estruturais, divergências culturais e transferências iconográficas transnacionais.")

abnt_paragraph(doc,
    "A unidade de análise é o item iconográfico — cada representação alegórica feminina em seu "
    "suporte material específico — e não o texto jurídico nem o sujeito individual. O corpus é, "
    "portanto, um corpus de imagens primárias, complementado por fontes documentais contextuais "
    "que permitem reconstituir as condições de produção, circulação e recepção de cada item. A "
    "dimensão empírica da pesquisa é fundamentalmente documental e arquivística: não se realizam "
    "entrevistas nem surveys, pois os sujeitos históricos não são acessíveis e a evidência "
    "pertinente reside nos objetos visuais e nos documentos que os enquadram institucionalmente.")

# ══════════════════════════════════════════════════════════════════════
# 3. MÉTODO CENTRAL: ICONOLOGIA JURÍDICA
# ══════════════════════════════════════════════════════════════════════
heading(doc, "3  Método Central: A Iconologia Jurídica")

abnt_paragraph(doc,
    "O método central desta tese é a iconologia jurídica — adaptação do método iconológico de Erwin "
    "Panofsky ao campo do direito. Na formulação original de Panofsky, elaborada em Meaning in the "
    "Visual Arts (1955), a análise de imagens avança por três níveis progressivos de profundidade "
    "interpretativa. O primeiro é o nível pré-iconográfico, que descreve o que é visualmente "
    "perceptível sem recorrer à interpretação convencional: formas, postura, gesto, composição "
    "espacial, suporte material, texto inscrito. O segundo é o nível iconográfico, que identifica "
    "o significado convencional dos motivos e símbolos reconhecíveis pela cultura visual "
    "compartilhada: a balança como Justiça, o barrete frígio como Liberdade, o capacete como "
    "potência bélica, a cornucópia como prosperidade. O terceiro é o nível iconológico, que "
    "interpreta o significado intrínseco ou sintomático da obra, revelando os valores, "
    "pressupostos e tensões da época em que foi produzida.")

abnt_paragraph(doc,
    "A adaptação jurídica deste método opera um deslocamento fundamental no terceiro nível: o que "
    "se interpreta não é apenas o 'espírito da época' no sentido burckardtiano, mas a função "
    "jurídico-política específica da alegoria — como ela legitima, estabiliza ou contesta a ordem "
    "estatal, como instrumentaliza o corpo feminino para fins de governo simbólico e como se inscreve "
    "em um regime iconocrático determinado. A iconologia jurídica, portanto, não trata as imagens "
    "como ilustrações do direito, mas como documentos jurídicos primários, dotados de eficácia "
    "normativa análoga à dos textos legais.")

abnt_paragraph(doc,
    "A operacionalização deste método sobre o corpus se dá em três etapas articuladas. Na primeira "
    "etapa — análise pré-iconográfica —, cada item é descrito em termos de suas propriedades visuais "
    "objetivas: número e postura das figuras, vestimenta e grau de cobertura corporal, atributos "
    "físicos carregados (espada, balança, barrete frígio, coroa, escudo, fasces, tocha, livro de leis), "
    "composição espacial, técnica e suporte, texto inscrito. Na segunda etapa — análise iconográfica "
    "—, identifica-se o motivo alegórico específico (Marianne, Britannia, Germania, Columbia, Justitia, "
    "A República), o código Iconclass primário (48C51 para alegorias femininas estatais), a tradição "
    "iconográfica citada e a Pathosformel warburguiana ativa — a fórmula gestual que porta, em si "
    "mesma, uma memória de emoção histórica (Nachleben). Na terceira etapa — análise iconológica —, "
    "interpreta-se a função jurídico-política da alegoria, classificando-a em um dos três regimes "
    "iconocráticos (FUNDACIONAL, NORMATIVO, MILITAR) e avaliando o grau de ENDURECIMENTO do corpo "
    "alegórico mediante os dez indicadores de purificação.")

# ══════════════════════════════════════════════════════════════════════
# 4. ENDURECIMENTO E REGIMES ICONOCRÁTICOS
# ══════════════════════════════════════════════════════════════════════
heading(doc, "4  Os Três Regimes Iconocráticos e o Conceito de ENDURECIMENTO")

abnt_paragraph(doc,
    "O conceito de iconocracia designa o regime pelo qual o Estado moderno utiliza imagens — "
    "especialmente imagens do corpo feminino alegorizado — como instrumento de produção, "
    "estabilização e contestação da ordem jurídica legítima. Esta tese propõe uma tipologia de "
    "três regimes iconocráticos que correspondem a fases históricas e funções políticas distintas.")

abnt_paragraph(doc,
    "O regime FUNDACIONAL caracteriza a fase de constituição dos Estados republicanos: o corpo "
    "alegórico está em seu estado mais vivo, dinâmico e exposto. A Marianne revolucionária de "
    "Delacroix — de peito nu, barrete frígio, bandeira erguida — é o paradigma deste regime. O "
    "corpo feminino funciona aqui como dispositivo sacrificial e fundador: ele encarna a ruptura "
    "com a ordem anterior e legitima a nova ordem pelo gesto e pela carne. O regime NORMATIVO "
    "caracteriza a fase de estabilização burocrática da república: a alegoria é domesticada, "
    "vestida, serializada e infinitamente reproduzível. A Semeuse de Oscar Roty, adotada nos "
    "selos e moedas francesas a partir de 1898, é o paradigma — corpo caminhante mas coberto, "
    "de perfil, sem dramatismo, perfeitamente reproduzível em milhões de exemplares idênticos. "
    "O regime MILITAR caracteriza a fase de mobilização imperial e bélica: o corpo alegórico "
    "ENDURECE ao máximo, armando-se, blindando-se, dominando territórios e sujeitos coloniais. "
    "A Britannia com tridente e capacete, sentada sobre o globo, é o seu paradigma.")

abnt_paragraph(doc,
    "O ENDURECIMENTO — conceito original desta tese — designa o processo de purificação progressiva "
    "pelo qual o corpo feminino alegorizado se transforma de corpo vivo em ícone estatal imóvel, "
    "intercambiável e metalizado. Trata-se de um processo mensurável a partir de dez indicadores "
    "analíticos, cada um avaliado em escala ordinal de zero a quatro pontos (0 = ausente; 4 = "
    "extremo): (1) desincorporação — redução do corpo inteiro ao busto, ao rosto ou ao símbolo; "
    "(2) rigidez postural — congelamento progressivo do movimento; (3) dessexualização — ocultação "
    "do corpo e remoção de marcadores de corporalidade feminina; (4) uniformização facial — "
    "genericidade crescente das feições; (5) heraldicização — integração em programa heráldico "
    "ou institucional; (6) enquadramento arquitetônico — emolduramento por bordas arquitetônicas "
    "ou ornamentais; (7) apagamento narrativo — remoção do contexto de ação; (8) monocromatização "
    "— redução da policromia ao monocromo ou ao metálico; (9) serialidade — reprodução industrial "
    "em massa; e (10) inscrição estatal — texto ou símbolos estatais inscritos sobre ou ao redor "
    "do corpo. O score de ENDURECIMENTO de cada item é calculado como média aritmética dos dez "
    "indicadores (escala 0,0–4,0), funcionando sempre como síntese interpretativa, nunca como "
    "substituição da análise qualitativa individual de cada indicador.")

abnt_paragraph(doc,
    "Esta tese propõe igualmente dois conceitos originais que operacionalizam a dimensão de gênero "
    "do ENDURECIMENTO. O Contrato Sexual Visual designa o mecanismo pelo qual o Estado moderno "
    "instrumentaliza o corpo feminino alegorizado para fins de legitimação jurídico-política: o corpo "
    "da mulher — abstrato, idealizado, purificado — é convertido em suporte da autoridade estatal "
    "precisamente porque é esvaziado de sua particularidade de sujeito. A Feminilidade de Estado "
    "designa a feminilidade como tecnologia de governo visual: não uma feminilidade qualquer, mas "
    "aquela construída, selecionada e estabilizada pelo aparato estatal para representar os atributos "
    "que o Estado deseja ver associados à sua própria imagem — pureza, imparcialidade, majestade, fecundidade.")

# ══════════════════════════════════════════════════════════════════════
# 5. CONSTITUIÇÃO DO CORPUS
# ══════════════════════════════════════════════════════════════════════
heading(doc, "5  Constituição do Corpus")

abnt_paragraph(doc,
    "O corpus da tese é composto por representações alegóricas femininas com função jurídico-política "
    "explícita, produzidas entre 1800 e 2000 em seis jurisdições nacionais: França (FR), Reino Unido "
    "(UK), Alemanha (DE), Estados Unidos (US), Bélgica (BE) e Brasil (BR). A seleção dessas "
    "jurisdições não é arbitrária: corresponde ao conjunto de países que, no período estudado, "
    "produziram as tradições iconográficas mais sistematizadas de alegoria feminina estatal — e que, "
    "ademais, mantiveram relações de influência e transferência recíproca, constituindo um sistema "
    "iconográfico transnacional cujas circulações a tese visa mapear.")

abnt_paragraph(doc,
    "A inclusão de um item no corpus exige o cumprimento simultâneo de cinco critérios cumulativos: "
    "(1) a figura representada deve ser feminina e alegórica, excluindo-se retratos de mulheres reais "
    "e figuras masculinas; (2) deve exercer função jurídico-política explícita — representando a "
    "República, a Justiça, a Nação, a Lei ou o Estado —, excluindo-se figuras meramente decorativas "
    "ou mitológicas sem conexão com a ordem jurídica; (3) deve ser datável entre 1800 e 2000; "
    "(4) deve pertencer a uma das seis jurisdições contempladas; e (5) deve estar inscrita em um "
    "dos suportes materiais aceitos: moeda circulante ou colonial, selo postal, monumento público "
    "ou escultura, arquitetura forense (inclusive decoração interior de tribunais — murais, vitrais, "
    "relevos de vestíbulo, pinturas de teto), estampa ou gravura, frontispício de codificação legal, "
    "papel-moeda e cartaz oficial ou de propaganda.")

abnt_paragraph(doc,
    "A estratégia de constituição do corpus é exaustiva dentro dos critérios estabelecidos — não "
    "por amostragem aleatória, mas por saturação informacional: a busca continua até que novas "
    "buscas sistemáticas nos acervos não produzam candidatos novos que satisfaçam os critérios "
    "de inclusão. Os acervos prioritários são aqueles com metadados verificáveis e imagens de alta "
    "resolução disponíveis em domínio público ou mediante licença: a Bibliothèque nationale de "
    "France (Gallica/BnF), o British Museum Collection Online, a Library of Congress Digital "
    "Collections, os portais de numismática comparada Numista e Colnect, a Biblioteca Nacional "
    "Digital do Brasil, o Musée des monnaies et médailles de Paris e a Bundesdruckerei. Para cada "
    "item selecionado, é registrada uma ficha-padrão contendo metadados completos — título, data, "
    "jurisdição, suporte, dimensões, acervo, URL ou referência de localização, direitos de "
    "reprodução — e a análise nos três níveis iconológicos.")

# ══════════════════════════════════════════════════════════════════════
# 6. TÉCNICAS DE PESQUISA
# ══════════════════════════════════════════════════════════════════════
heading(doc, "6  Técnicas de Pesquisa")

abnt_paragraph(doc,
    "A tese mobiliza três técnicas complementares que operam em sequência e de modo mutuamente "
    "informado. A primeira é a pesquisa documental e arquivística, que consiste na localização, "
    "recuperação e catalogação sistemática de fontes primárias nos acervos identificados. Esta "
    "técnica opera em dois planos simultâneos: o levantamento iconográfico — identificação das "
    "alegorias femininas nos suportes materiais aceitos — e o levantamento documental contextual "
    "— leis, decretos, relatórios de comissões parlamentares, concursos de cunhagem, correspondência "
    "oficial e registros administrativos que permitam reconstituir as condições de produção de cada "
    "item. O levantamento contextual é metodologicamente indispensável porque evita interpretações "
    "anacrônicas: saber que um concurso de Marianne foi politicamente disputado, ou que uma alegoria "
    "foi recusada pelo Tesouro por 'indecência', transforma radicalmente a interpretação iconológica "
    "do item resultante.")

abnt_paragraph(doc,
    "A segunda técnica é a análise de conteúdo visual, entendida à luz da proposta metodológica de "
    "Laurence Bardin e adaptada às especificidades do documento iconográfico. As categorias analíticas "
    "são estabelecidas a priori com base no marco teórico da pesquisa e operacionalizadas pelo "
    "protocolo dos dez indicadores de ENDURECIMENTO: cada indicador funciona como uma categoria de "
    "análise com definição operacional explícita, escala de pontuação e exemplos-âncora calibrados "
    "por subgrupo de regime iconocrático. A análise é realizada em dupla rodada independente — com "
    "intervalo mínimo de quinze dias entre as rodadas — para controle da consistência interna. As "
    "divergências entre rodadas são reconciliadas por reflexão argumentada e documentadas no "
    "anexo metodológico.")

abnt_paragraph(doc,
    "A terceira técnica é a comparação sistemática entre jurisdições. A comparação opera em dois "
    "níveis articulados. No primeiro nível, comparam-se os perfis médios de ENDURECIMENTO entre "
    "jurisdições e entre regimes iconocráticos, identificando tendências convergentes e divergentes. "
    "No segundo nível — o mais relevante para a tese —, analisam-se as transferências iconográficas "
    "transnacionais: como a Marianne francesa influenciou a República brasileira; como a Britannia "
    "colonial foi adaptada para as moedas do Império destinadas à Ásia; como a Germania "
    "wilhelminiana dialogou com a Minerva dos tribunais belgas. É neste segundo nível que a noção "
    "warburguiana de Zwischenraum — o espaço entre imagens de jurisdições distintas, prenhe de "
    "memória gestual e transferência formal — adquire sua plena potência analítica para a tese.")

# ── SEÇÃO 7 ─────────────────────────────────────────────────────────────────
heading(doc, "7  Validade, Confiabilidade e Consistência")

abnt_paragraph(doc,
    "A questão da validade em pesquisa histórico-iconográfica não se resolve pelos "
    "critérios estatísticos da replicabilidade experimental, mas exige um conjunto "
    "articulado de dispositivos que assegurem, simultaneamente, a consistência "
    "interna do método e a rastreabilidade intersubjetiva das interpretações. Para "
    "a presente tese, quatro dispositivos foram concebidos de forma integrada. O "
    "primeiro é a triangulação metodológica: cada peça do corpus é examinada por "
    "ao menos três vias analíticas independentes — a análise iconológica em três "
    "níveis, a pontuação dos dez indicadores de ENDURECIMENTO e a localização da "
    "peça no mapa comparativo do Zwischenraum — de modo que nenhuma conclusão "
    "interpretativa repousa sobre uma única operação analítica."
)

abnt_paragraph(doc,
    "O segundo dispositivo é a operacionalização explícita e pública dos conceitos. "
    "Os dez indicadores de ENDURECIMENTO são definidos operacionalmente no protocolo "
    "IconoCode com descrição verbal, exemplos positivos e negativos, e critérios de "
    "pontuação (escala ordinal 0–4) suficientemente precisos para permitir a "
    "aplicação independente por outros pesquisadores. Essa transparência definitória "
    "constitui a base da validade de constructo: o leitor pode contestar a "
    "pertinência da escala, mas não a arbitrariedade de sua aplicação ao corpus."
)

abnt_paragraph(doc,
    "O terceiro dispositivo é a dupla codificação amostral: uma subamostra de "
    "aproximadamente 20% do corpus (selecionada por sorteio estratificado por país "
    "e suporte) será codificada de forma independente por um segundo pesquisador "
    "treinado no protocolo IconoCode, calculando-se o índice de concordância "
    "inter-avaliador (Kappa de Cohen). Valores de κ ≥ 0,70 são tomados como limiar "
    "de confiabilidade aceitável; divergências sistemáticas serão registradas e "
    "utilizadas para refinamento dos descritores operacionais. O quarto dispositivo "
    "é o peer debriefing estruturado: capítulos analíticos serão submetidos, em "
    "versão preliminar, a especialistas externos em história da arte jurídica e em "
    "estudos de gênero, cujas objeções serão respondidas ou incorporadas antes da "
    "versão final, assegurando validade comunicativa e ancoragem disciplinar."
)

abnt_paragraph(doc,
    "A validade externa da tese é de natureza analítica, não estatística: o objetivo "
    "não é generalizar para uma população de imagens, mas testar a transferibilidade "
    "do modelo de ENDURECIMENTO para contextos culturais comparáveis. A estratégia "
    "de saturação teórica — encerramento da coleta quando novos itens deixam de "
    "acrescentar variação nos scores dos indicadores — fornece o critério de "
    "suficiência do corpus sem recorrer a cálculos amostrais probabilísticos. "
    "O conjunto dessas medidas confere à pesquisa o grau de rigor metodológico "
    "compatível com as exigências epistemológicas de uma tese histórico-interpretativa "
    "de doutoramento."
)

# ── SEÇÃO 8 ─────────────────────────────────────────────────────────────────
heading(doc, "8  Vieses e Limitações Declaradas")

abnt_paragraph(doc,
    "A declaração antecipada de limitações não enfraquece a tese; ao contrário, "
    "é condição de honestidade intelectual e de utilidade crítica para pesquisas "
    "subsequentes. A primeira limitação de ordem estrutural é o viés digital de "
    "arquivo: o corpus é constituído exclusivamente por peças disponíveis em acervos "
    "digitalizados publicamente acessíveis (Gallica/BnF, Europeana, Library of "
    "Congress, Brasiliana Digital, British Museum Online, Numista, Colnect). Essa "
    "escolha metodológica deliberada garante rastreabilidade e reprodutibilidade, "
    "mas introduz viés de sobrerrepresentação de itens provenientes de acervos "
    "europeus bem digitalizados — em especial franceses — e de subrepresentação "
    "de materiais brasileiros e belgas ainda em processo de digitalização. Os "
    "resultados comparativos devem ser lidos à luz desse desequilíbrio."
)

abnt_paragraph(doc,
    "A segunda limitação é o viés linguístico de cobertura bibliográfica: a revisão "
    "de literatura foi conduzida prioritariamente em francês, inglês e português, "
    "deixando potencialmente de lado contribuições relevantes em alemão, neerlandês "
    "e espanhol, idiomas nos quais parte substancial da historiografia sobre Germania, "
    "La Belgique e alegorias latino-americanas foi produzida. Esforços de ampliação "
    "linguística foram feitos onde possível, mas a limitação permanece e é declarada."
)

abnt_paragraph(doc,
    "A terceira limitação é inerente ao método interpretativo: a pontuação dos dez "
    "indicadores de ENDURECIMENTO depende, em última instância, do juízo informado "
    "da pesquisadora, cuja trajetória intelectual (formada na interface entre direito "
    "e história da arte) constitui simultaneamente o ativo e o limite da análise. "
    "O protocolo de dupla codificação reduz, mas não elimina, a dimensão subjetiva "
    "da leitura iconológica. A quarta limitação é de alcance geográfico-temporal: "
    "a tese cobre seis jurisdições entre 1800 e 2000, recorte que exclui "
    "deliberadamente o período contemporâneo (pós-2000) e jurisdições do Sul Global "
    "não colonizadoras, nos quais a alegoria feminina opera sob lógicas que "
    "merecem investigação autônoma. Essas exclusões são teóricas, não empíricas: "
    "definem o escopo legítimo de generalização analítica dos resultados desta tese."
)

# ── REFERÊNCIAS ──────────────────────────────────────────────────────────────
heading(doc, "Referências")

refs = [
    ("AGULHON, Maurice. ",
     "Marianne au combat: l'imagerie et la symbolique républicaines de 1789 à 1880. "
     "Paris: Flammarion, 1979."),
    ("BARDIN, Laurence. ",
     "Análise de conteúdo. Tradução de Luís Antero Reto e Augusto Pinheiro. "
     "São Paulo: Edições 70, 2011."),
    ("DAL RI JR., Arno. ",
     "História do direito internacional: comércio e moeda, cidadania e nacionalidade. "
     "Florianópolis: Fundação Boiteux, 2004."),
    ("HUYGEBAERT, Stefan et al. ",
     "The art of law: three centuries of justice depicted. Tielt: Lannoo, 2016."),
    ("KOCHER, Gernot; SCHOTT, Clausdieter (Hrsg.). ",
     "Bilder aus der deutschen Rechtsgeschichte: von den Anfängen bis zur Gegenwart. "
     "München: C. H. Beck, 1992."),
    ("PANOFSKY, Erwin. ",
     "Estudos de iconologia: temas humanísticos na arte do Renascimento. "
     "Tradução de Olinda Braga de Sousa. Lisboa: Estampa, 1986."),
    ("PATEMAN, Carole. ",
     "The sexual contract. Cambridge: Polity Press, 1988."),
    ("WARNER, Marina. ",
     "Monuments and maidens: the allegory of the female form. "
     "London: Weidenfeld and Nicolson, 1985."),
]

for bold_part, normal_part in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent  = Cm(0)
    pf.left_indent        = Cm(1.25)
    pf.space_before       = Pt(0)
    pf.space_after        = Pt(6)
    pf.line_spacing_rule  = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing       = 1.5
    rb = p.add_run(bold_part)
    rb.bold = True
    rb.font.name = "Times New Roman"
    rb.font.size = Pt(12)
    rn = p.add_run(normal_part)
    rn.font.name = "Times New Roman"
    rn.font.size = Pt(12)

# ── SALVAR ───────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print("Documento salvo em:", OUTPUT)
