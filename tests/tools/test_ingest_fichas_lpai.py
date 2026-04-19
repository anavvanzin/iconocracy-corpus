"""Tests for ingest_fichas_lpai.py — T4 staging parser."""

from __future__ import annotations

import io
import json
import sys
import zipfile
from pathlib import Path

import pytest

# Add tools/scripts to path
REPO_ROOT = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(REPO_ROOT / "tools" / "scripts"))

import ingest_fichas_lpai as ingest  # noqa: E402


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------
FIELD_LABELS = [
    ("═══  LPAI v2 RECORD  ═══", "═══  LPAI v2 RECORD  ═══"),
    ("ID", "BR-SCOUT-001"),
    ("Título", "Alegoria da República Brasileira — Revista Illustrada"),
    ("Autoria", "Angelo Agostini (1843–1910)"),
    ("Data", "c. 1889"),
    ("Fonte", "Hemeroteca Digital Brasileira / Wikimedia Commons"),
    ("URL", "https://commons.wikimedia.org/wiki/File:Republica_no_brasil.jpg"),
    (
        "URL download",
        "https://upload.wikimedia.org/wikipedia/commons/0/0a/Republica_no_brasil.jpg",
    ),
    ("Suporte", "Litografia em periódico"),
    ("Dimensões", "455 × 681 px"),
    ("───────────────────", "───────────────────"),
    ("CÓDIGO LPAI v2", "312.4+210 / 211.1+321.2 / FUND.ins"),
    ("CLASSE", "312.4 (Efígie da República BR) + 210 (Libertas)"),
    ("ATRIBUTOS", "211.1 (barrete frígio), 321.2 (bandeira nacional)"),
    ("ICONCLASS", "44B11 + 44B1 — mapeamento parcial"),
    ("REGIME", "FUNDACIONAL"),
    ("MODO", "Insurgente"),
    ("NOTA ANALÍTICA", "Primeira representação icônica pós-Proclamação."),
    ("REF. ABNT", "AGOSTINI, Angelo. Alegoria. 1889."),
]


def _synth_docx(tmp_path: Path, num_fichas: int = 15) -> Path:
    """Build a minimal DOCX file with `num_fichas` ficha tables using python-docx."""
    from docx import Document

    doc = Document()
    # Add some non-ficha content first
    doc.add_paragraph("INTRODUCTORY PARAGRAPH")
    for i in range(num_fichas):
        t = doc.add_table(rows=len(FIELD_LABELS), cols=2)
        for r, (label, value) in enumerate(FIELD_LABELS):
            if label == "ID":
                value = f"BR-SCOUT-{i + 1:03d}" if i < 7 else f"FR-SCOUT-{i - 6:03d}"
            if label == "URL":
                value = f"https://example.org/ficha/{i + 1}"
            if label == "Título":
                value = f"Ficha de teste {i + 1}"
            t.rows[r].cells[0].text = label
            t.rows[r].cells[1].text = value
    # Non-ficha table (footer-like)
    extra = doc.add_table(rows=2, cols=5)
    extra.rows[0].cells[0].text = "Arquivo"
    extra.rows[0].cells[1].text = "Coleção"

    out = tmp_path / "synth.docx"
    doc.save(str(out))
    return out


@pytest.fixture
def synth_docx(tmp_path):
    return _synth_docx(tmp_path, num_fichas=15)


# ---------------------------------------------------------------------------
# Parser tests
# ---------------------------------------------------------------------------
def test_parser_finds_fifteen_fichas(synth_docx):
    fichas = ingest.parse_docx(synth_docx)
    assert len(fichas) == 15


def test_parser_extracts_required_fields(synth_docx):
    fichas = ingest.parse_docx(synth_docx)
    f = fichas[0]
    assert f["id"] == "BR-SCOUT-001"
    assert f["title"].startswith("Ficha de teste")
    assert f["url"].startswith("https://example.org/ficha/")
    assert f["regime"] == "FUNDACIONAL"
    assert f["modo"] == "Insurgente"
    assert "ficha" in f["title"].lower()


def test_separator_rows_are_skipped(synth_docx):
    fichas = ingest.parse_docx(synth_docx)
    for f in fichas:
        assert all(not k.startswith("───") for k in f)
        assert "═══  LPAI v2 RECORD  ═══" not in f


# ---------------------------------------------------------------------------
# Dedup tests
# ---------------------------------------------------------------------------
def test_dedup_detects_url_match(tmp_path):
    corpus = tmp_path / "corpus.json"
    records = tmp_path / "records.jsonl"
    corpus.write_text(
        json.dumps(
            [
                {
                    "id": "XX-001",
                    "title": "some prior item",
                    "url": "https://example.org/ficha/1",
                }
            ]
        ),
        encoding="utf-8",
    )
    records.write_text("", encoding="utf-8")
    idx = ingest.DedupIndex.from_files(corpus, records)
    status, match = idx.classify("BR-SCOUT-001", "https://example.org/ficha/1", "New title")
    assert status == "PARTIAL"
    assert match == "XX-001"


def test_dedup_detects_title_match(tmp_path):
    corpus = tmp_path / "corpus.json"
    records = tmp_path / "records.jsonl"
    corpus.write_text(
        json.dumps(
            [{"id": "XX-042", "title": "Alegoria da República — Teste"}]
        ),
        encoding="utf-8",
    )
    records.write_text("", encoding="utf-8")
    idx = ingest.DedupIndex.from_files(corpus, records)
    status, match = idx.classify(
        "BR-SCOUT-999", "https://brand-new/url", "alegoria da republica teste"
    )
    # title matches after normalization
    assert status == "PARTIAL"
    assert match == "XX-042"


def test_dedup_strong_match_two_signals(tmp_path):
    corpus = tmp_path / "corpus.json"
    records = tmp_path / "records.jsonl"
    corpus.write_text(
        json.dumps(
            [
                {
                    "id": "XX-007",
                    "title": "Exact Title",
                    "url": "https://u/1",
                }
            ]
        ),
        encoding="utf-8",
    )
    records.write_text("", encoding="utf-8")
    idx = ingest.DedupIndex.from_files(corpus, records)
    status, match = idx.classify("BR-SCOUT-001", "https://u/1", "Exact Title")
    assert status == "MATCHES"
    assert match == "XX-007"


def test_dedup_new_item(tmp_path):
    corpus = tmp_path / "corpus.json"
    corpus.write_text("[]", encoding="utf-8")
    records = tmp_path / "records.jsonl"
    records.write_text("", encoding="utf-8")
    idx = ingest.DedupIndex.from_files(corpus, records)
    status, match = idx.classify("BR-SCOUT-001", "https://new/1", "new title")
    assert status == "NEW"
    assert match is None


# ---------------------------------------------------------------------------
# Schema validation
# ---------------------------------------------------------------------------
def test_record_passes_master_schema(synth_docx):
    fichas = ingest.parse_docx(synth_docx)
    rec = ingest.build_staging_record(
        fichas[0], batch_id="00000000-0000-4000-8000-lpaiv2scout0001", now_iso="2026-04-19T12:00:00Z"
    )
    errors = ingest.validate_record(rec)
    assert errors == [], f"unexpected schema errors: {errors}"


# ---------------------------------------------------------------------------
# End-to-end run()
# ---------------------------------------------------------------------------
def test_dry_run_writes_nothing(synth_docx, tmp_path, monkeypatch):
    stage = tmp_path / "stage"
    result = ingest.run(synth_docx, stage, dry_run=True, skip_images=True)
    assert not stage.exists()
    assert result["dry_run"] is True
    assert len(result["fichas"]) == 15
    assert result["written_files"] == []


def test_wet_run_writes_jsonl_and_drafts(synth_docx, tmp_path):
    stage = tmp_path / "stage"
    result = ingest.run(synth_docx, stage, skip_images=True)
    jsonl = Path(result["jsonl_path"])
    assert jsonl.exists()
    lines = jsonl.read_text(encoding="utf-8").strip().splitlines()
    assert len(lines) == 15
    # Each line parses as JSON with required top-level keys
    for line in lines:
        rec = json.loads(line)
        assert rec["master_record_version"] == "1.0"
        assert rec["input"]["input_url"].startswith("https://")
    drafts_dir = Path(result["drafts_dir"])
    drafts = list(drafts_dir.glob("*.md"))
    assert len(drafts) == 15


def test_skip_images_does_not_create_images_dir(synth_docx, tmp_path):
    stage = tmp_path / "stage"
    result = ingest.run(synth_docx, stage, skip_images=True)
    images_dir = Path(result["images_dir"])
    assert not images_dir.exists()


def test_extract_images_creates_images_dir_even_when_docx_has_none(synth_docx, tmp_path):
    stage = tmp_path / "stage"
    result = ingest.run(synth_docx, stage, skip_images=False)
    images_dir = Path(result["images_dir"])
    assert images_dir.exists()
    # synth docx has no inline images
    assert result["extracted_images"] == []


def test_extract_images_writes_files_when_docx_has_media(tmp_path):
    # Build a docx with a fake inline media entry by post-processing the zip.
    from docx import Document

    doc = Document()
    t = doc.add_table(rows=len(FIELD_LABELS), cols=2)
    for r, (label, value) in enumerate(FIELD_LABELS):
        t.rows[r].cells[0].text = label
        t.rows[r].cells[1].text = value
    src = tmp_path / "src.docx"
    doc.save(str(src))

    # Inject a fake PNG into word/media/
    patched = tmp_path / "with_media.docx"
    fake_png = b"\x89PNG\r\n\x1a\n" + b"FAKE" * 4
    with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(patched, "w") as zout:
        for item in zin.namelist():
            zout.writestr(item, zin.read(item))
        zout.writestr("word/media/image1.png", fake_png)

    stage = tmp_path / "stage"
    result = ingest.run(patched, stage, skip_images=False)
    images_dir = Path(result["images_dir"])
    assert (images_dir / "image1.png").exists()
    assert "image1.png" in result["extracted_images"]


def test_exit_code_two_when_validation_fails(synth_docx, tmp_path, monkeypatch):
    """If any record fails schema validation, main() returns 2 but still writes staging."""
    stage = tmp_path / "stage"

    # Force validation failure by patching validate_record to always report one error
    monkeypatch.setattr(
        ingest,
        "validate_record",
        lambda record: ["root: forced test error"],
    )
    rc = ingest.main(
        [
            "--input",
            str(synth_docx),
            "--stage-dir",
            str(stage),
            "--skip-images",
        ]
    )
    assert rc == 2
    # Staging still written
    assert (stage / ingest.STAGING_JSONL_NAME).exists()


def test_exit_code_zero_on_success(synth_docx, tmp_path):
    stage = tmp_path / "stage"
    rc = ingest.main(
        [
            "--input",
            str(synth_docx),
            "--stage-dir",
            str(stage),
            "--skip-images",
        ]
    )
    assert rc == 0
